"""Unit tests for the PDF→Word pre-flight and fallback conversion paths.

Uses real PDFs built in-test with PyMuPDF/pikepdf — the failure modes under
test (encryption, damaged structure) live in the file format itself and
cannot be faked meaningfully.
"""

from __future__ import annotations

import pytest

from app.exceptions.jobs import ProcessingError
from app.utils.pdf_convert import (
    _fallback_pdf_to_docx,
    apply_rtl_direction,
    pdf_to_docx,
    prepare_pdf_for_conversion,
)

#: Arabic with an embedded number, percentage and Latin word — the mixed run
#: that renders wrongly when the paragraph has no base direction.
ARABIC = "مرحبا بالعالم 2026 والنسبة %100 (اختبار) MARKET"


@pytest.fixture
def text_pdf(tmp_path):
    """A small, clean, unencrypted PDF with real text content."""
    import fitz

    path = tmp_path / "clean.pdf"
    doc = fitz.open()
    for number in range(2):
        page = doc.new_page()
        page.insert_text((72, 100), f"Hello page {number + 1}", fontsize=14)
    doc.save(str(path))
    doc.close()
    return path


def _encrypt(source, target, *, owner: str, user: str) -> None:
    import pikepdf

    with pikepdf.open(source) as pdf:
        pdf.save(
            target,
            encryption=pikepdf.Encryption(owner=owner, user=user, R=6),
        )


class TestPreparePdfForConversion:
    def test_clean_pdf_returned_untouched(self, text_pdf, tmp_path):
        workspace = tmp_path / "ws"
        workspace.mkdir()
        assert prepare_pdf_for_conversion(text_pdf, workspace) == text_pdf
        assert not list(workspace.iterdir())  # no needless copies

    def test_owner_password_pdf_passes_through(self, text_pdf, tmp_path):
        """Owner-password PDFs auto-authenticate in PyMuPDF (pdf2docx reads
        them fine), so the fast path must not copy them needlessly."""
        locked = tmp_path / "owner-locked.pdf"
        _encrypt(text_pdf, locked, owner="secret", user="")
        workspace = tmp_path / "ws"
        workspace.mkdir()

        assert prepare_pdf_for_conversion(locked, workspace) == locked

    def test_unreadable_for_fitz_is_repaired_via_pikepdf(
        self, text_pdf, tmp_path, monkeypatch
    ):
        """Files PyMuPDF cannot parse get a pikepdf structural rewrite."""
        import fitz

        def refuse(*_args, **_kwargs):
            raise RuntimeError("cannot open broken xref")

        monkeypatch.setattr(fitz, "open", refuse)
        workspace = tmp_path / "ws"
        workspace.mkdir()

        prepared = prepare_pdf_for_conversion(text_pdf, workspace)

        assert prepared != text_pdf
        assert prepared.parent == workspace
        assert prepared.stat().st_size > 0

    def test_user_password_pdf_raises_actionable_error(self, text_pdf, tmp_path):
        locked = tmp_path / "user-locked.pdf"
        _encrypt(text_pdf, locked, owner="secret", user="secret")
        workspace = tmp_path / "ws"
        workspace.mkdir()

        with pytest.raises(ProcessingError, match="password-protected"):
            prepare_pdf_for_conversion(
                locked, workspace, display_name="report.pdf"
            )

    def test_garbage_raises_damaged_error(self, tmp_path):
        garbage = tmp_path / "junk.pdf"
        garbage.write_bytes(b"%PDF-1.7 not really a pdf at all")
        workspace = tmp_path / "ws"
        workspace.mkdir()

        with pytest.raises(ProcessingError, match=r"damaged|not a valid"):
            prepare_pdf_for_conversion(garbage, workspace)


class TestPdfToDocx:
    def test_converts_clean_pdf(self, text_pdf, tmp_path):
        output = tmp_path / "out.docx"
        assert pdf_to_docx(text_pdf, output) == output
        assert output.stat().st_size > 0

    def test_owner_locked_pdf_converts_after_preflight(self, text_pdf, tmp_path):
        """The full task-side chain: normalize, then convert."""
        locked = tmp_path / "owner-locked.pdf"
        _encrypt(text_pdf, locked, owner="secret", user="")
        workspace = tmp_path / "ws"
        workspace.mkdir()

        prepared = prepare_pdf_for_conversion(locked, workspace)
        output = tmp_path / "out.docx"
        assert pdf_to_docx(prepared, output).stat().st_size > 0

    def test_fallback_extracts_text(self, text_pdf, tmp_path):
        from docx import Document

        output = tmp_path / "fallback.docx"
        _fallback_pdf_to_docx(text_pdf, output, first_page=None, last_page=None)

        text = "\n".join(p.text for p in Document(str(output)).paragraphs)
        assert "Hello page 1" in text
        assert "Hello page 2" in text

    def test_fallback_respects_page_range(self, text_pdf, tmp_path):
        from docx import Document

        output = tmp_path / "fallback-range.docx"
        _fallback_pdf_to_docx(text_pdf, output, first_page=2, last_page=2)

        text = "\n".join(p.text for p in Document(str(output)).paragraphs)
        assert "Hello page 1" not in text
        assert "Hello page 2" in text


def _body_xml(path) -> str:
    import zipfile

    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _write_docx(path, *paragraphs, in_table: bool = False):
    from docx import Document

    document = Document()
    if in_table:
        table = document.add_table(rows=len(paragraphs), cols=1)
        for row, text in zip(table.rows, paragraphs):
            row.cells[0].paragraphs[0].add_run(text)
    else:
        for text in paragraphs:
            document.add_paragraph().add_run(text)
    document.save(str(path))
    return path


class TestRtlDirection:
    """pdf2docx reproduces glyph positions but emits no direction markup, so
    Arabic/Hebrew lands in LTR paragraphs and Word bidi-reorders it wrongly."""

    def test_arabic_paragraph_gets_bidi_and_rtl_runs(self, tmp_path):
        path = _write_docx(tmp_path / "ar.docx", ARABIC)

        assert apply_rtl_direction(path) == 1

        xml = _body_xml(path)
        assert "<w:bidi" in xml  # paragraph base direction
        assert "<w:rtl" in xml  # run direction

    def test_latin_document_is_left_untouched(self, tmp_path):
        path = _write_docx(tmp_path / "en.docx", "Hello world 2026 (test)")

        assert apply_rtl_direction(path) == 0
        assert "<w:bidi" not in _body_xml(path)

    def test_one_arabic_word_does_not_flip_an_english_paragraph(self, tmp_path):
        path = _write_docx(
            tmp_path / "mixed.docx",
            "The Arabic word for peace is سلام and it is written thus.",
        )

        assert apply_rtl_direction(path) == 0

    def test_paragraphs_inside_tables_are_marked(self, tmp_path):
        """Most pdf2docx output text lives in cells, not body paragraphs."""
        path = _write_docx(tmp_path / "tbl.docx", ARABIC, ARABIC, in_table=True)

        assert apply_rtl_direction(path) == 2

    def test_rerunning_does_not_duplicate_markup(self, tmp_path):
        path = _write_docx(tmp_path / "twice.docx", ARABIC)

        apply_rtl_direction(path)
        first = _body_xml(path).count("<w:bidi")
        apply_rtl_direction(path)

        assert _body_xml(path).count("<w:bidi") == first

    def test_property_order_is_normalised(self, tmp_path):
        """``w:pPr`` children are a schema sequence; pdf2docx emits some out
        of order, which puts indentation/justification at risk of being
        ignored. The pass must reseat them without dropping any."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        path = tmp_path / "unordered.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run(ARABIC)
        pPr = paragraph._p.get_or_add_pPr()
        for tag in ("w:jc", "w:ind", "w:spacing", "w:widowControl"):
            pPr.append(OxmlElement(tag))  # deliberately reversed
        document.save(str(path))

        apply_rtl_direction(path)

        reopened = Document(str(path))
        children = [
            qn(f"w:{t}")
            for t in ("widowControl", "bidi", "spacing", "ind", "jc")
        ]
        actual = [c.tag for c in reopened.paragraphs[0]._p.get_or_add_pPr()]
        assert actual == children

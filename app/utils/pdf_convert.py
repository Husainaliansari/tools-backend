"""PDF → Word conversion (pdf2docx + pikepdf pre-flight).

pdf2docx reconstructs editable DOCX from PDF layout analysis (tables,
paragraphs, images) — substantially better output than LibreOffice's PDF
import for text documents. The import is deferred: the library (PyMuPDF
backed) is heavy and only workers need it.

Reliability comes from bracketing pdf2docx with two safety nets:

* :func:`prepare_pdf_for_conversion` — pdf2docx refuses any encrypted PDF
  (even owner-password files every viewer opens) and trips over mildly
  corrupt xref tables. Such files are routed through pikepdf, which repairs
  structure and strips empty-password encryption; only a real user password
  is surfaced as an error.
* a PyMuPDF fallback — when pdf2docx's layout parser crashes on an exotic
  page, the file is re-converted text-block-by-text-block (page images for
  pages without text) instead of failing the job.

Conversion is CPU-bound pure Python, so threads win nothing (GIL); the
speed-ups here are process-based and used only where child processes can be
spawned (see :func:`can_spawn_processes`):

* several inputs — the caller converts files concurrently in a process pool;
* one large input — pdf2docx's own ``multi_processing`` splits pages across
  processes (:func:`page_parallel_workers` decides when it pays off).
"""

from __future__ import annotations

import io
import logging
import multiprocessing
import os
import re
from pathlib import Path

from app.exceptions.jobs import ProcessingError
from app.logging import get_logger

logger = get_logger(__name__)

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

#: Below this page count the per-process spawn/import overhead outweighs the
#: parallel parsing gain, so the conversion stays single-process.
_MIN_PAGES_FOR_PARALLEL = 8

#: Upper bound on worker processes — pdf2docx workers are memory-hungry and
#: page-splitting shows diminishing returns beyond this.
MAX_CONVERT_WORKERS = 4


def can_spawn_processes() -> bool:
    """Whether this process may create child processes.

    Daemonic processes (e.g. Celery prefork workers) cannot — parallel
    conversion silently degrades to single-process there.
    """
    try:
        return not multiprocessing.current_process().daemon
    except Exception:  # pragma: no cover - defensive
        return False


def page_parallel_workers(input_path: Path) -> int:
    """Workers for pdf2docx page-level multi-processing (0 = stay serial)."""
    if not can_spawn_processes():
        return 0
    try:
        import fitz

        with fitz.open(str(input_path)) as doc:
            pages = doc.page_count
    except Exception:
        return 0  # unreadable here — let the converter surface the real error
    if pages < _MIN_PAGES_FOR_PARALLEL:
        return 0
    return max(2, min(MAX_CONVERT_WORKERS, os.cpu_count() or 1))


def prepare_pdf_for_conversion(
    input_path: Path,
    workspace: Path,
    *,
    display_name: str | None = None,
) -> Path:
    """Return a PDF path pdf2docx can safely consume.

    Fast path (the overwhelming majority): a cleanly readable, unencrypted
    file is returned untouched. Encrypted files (pdf2docx rejects them all,
    including owner-password documents that open fine in any viewer) and
    files PyMuPDF cannot parse are rewritten through pikepdf, which repairs
    xref/structure damage and drops empty-password encryption. Only a PDF
    needing a real password is a hard error — with an actionable message.
    """
    name = display_name or input_path.name
    needs_normalize = False
    try:
        import fitz

        with fitz.open(str(input_path)) as doc:
            was_protected = doc.needs_pass
            if was_protected and not doc.authenticate(""):
                raise ProcessingError(
                    f"'{name}' is password-protected. Use the Unlock PDF tool "
                    "to remove the password, then convert it."
                )
            # authenticate() flips is_encrypted/needs_pass to False on success,
            # so the original protected state must be captured before calling it.
            needs_normalize = bool(doc.is_encrypted or was_protected)
            if not needs_normalize and doc.page_count == 0:
                raise ProcessingError(f"'{name}' contains no pages.")
    except ProcessingError:
        raise
    except Exception:
        # Unreadable for PyMuPDF — let pikepdf try a structural repair.
        needs_normalize = True

    if not needs_normalize:
        return input_path

    try:
        import pikepdf
    except ImportError:  # pragma: no cover - deployment problem
        return input_path  # degrade to the old behaviour rather than block

    normalized = workspace / f"{input_path.stem}-normalized.pdf"
    try:
        with pikepdf.open(input_path) as pdf:
            pdf.save(normalized)
    except pikepdf.PasswordError as exc:
        raise ProcessingError(
            f"'{name}' is password-protected. Use the Unlock PDF tool to "
            "remove the password, then convert it."
        ) from exc
    except Exception as exc:
        raise ProcessingError(
            f"'{name}' could not be read — the file appears to be damaged "
            "or is not a valid PDF."
        ) from exc
    logger.info("pdf_normalized_for_conversion", input=name)
    return normalized


#: Right-to-left scripts: Hebrew, Arabic, Syriac, Thaana, N'Ko, Samaritan,
#: Mandaic, Arabic Extended-A and the Arabic presentation-form blocks.
_RTL_CHARS = re.compile(
    "[֐-׿؀-ۿ܀-ݏހ-޿߀-߿"
    "ࠀ-࡟ࢠ-ࣿיִ-﷿ﹰ-﻿]"
)

#: A paragraph flips to RTL only when RTL characters dominate its letters —
#: one Arabic word quoted inside an English sentence must not flip it.
_RTL_DOMINANCE = 0.4

#: ``w:pPr`` / ``w:rPr`` children are an ordered *sequence* in OOXML — the
#: schema fixes their positions, and a property in the wrong slot is liable to
#: be ignored by the consumer. pdf2docx emits several out of order (notably
#: ``w:widowControl`` after ``w:autoSpace*`` and ``w:w`` before ``w:rFonts``),
#: which puts indentation/justification at risk, so the whole sequence is
#: re-sorted on the way out.
_PPR_ORDER = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
    "pPrChange",
)
_RPR_ORDER = (
    "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
    "dstrike", "outline", "shadow", "emboss", "imprint", "noProof",
    "snapToGrid", "vanish", "webHidden", "color", "spacing", "w", "kern",
    "position", "sz", "szCs", "highlight", "u", "effect", "bdr", "shd",
    "fitText", "vertAlign", "rtl", "cs", "em", "lang", "eastAsianLayout",
    "specVanish", "oMath",
)

#: Tags that follow the two elements inserted below, for
#: ``insert_element_before`` to land them in the right slot.
_PPR_AFTER_BIDI = tuple(
    f"w:{tag}" for tag in _PPR_ORDER[_PPR_ORDER.index("bidi") + 1:]
)
_RPR_AFTER_RTL = tuple(
    f"w:{tag}" for tag in _RPR_ORDER[_RPR_ORDER.index("rtl") + 1:]
)


def _iter_paragraphs(document):
    """Every paragraph in the document body, in order, each exactly once.

    Walking the XML beats recursing through ``tables``/``cells``: it reaches
    paragraphs at any nesting depth, and a merged cell — which the cell API
    repeats once per grid position it spans — is visited a single time.
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    for element in document.element.body.iter(qn("w:p")):
        yield Paragraph(element, document)


def _normalize_property_order(document) -> int:
    """Re-sort ``w:pPr``/``w:rPr`` children into schema order; returns fixes.

    Sorting is stable and only reorders known tags, so unrecognised extensions
    keep their relative position and nothing is dropped.
    """
    from docx.oxml.ns import qn

    ranks = {
        qn("w:pPr"): {qn(f"w:{t}"): i for i, t in enumerate(_PPR_ORDER)},
        qn("w:rPr"): {qn(f"w:{t}"): i for i, t in enumerate(_RPR_ORDER)},
    }
    fixed = 0
    for container_tag, rank in ranks.items():
        for element in document.element.body.iter(container_tag):
            children = list(element)
            if len(children) < 2:
                continue
            # Unknown tags stay put by ranking them where they already sit.
            keyed = [
                (rank.get(child.tag, rank.get(children[i - 1].tag, -1) if i else -1), i, child)
                for i, child in enumerate(children)
            ]
            ordered = [child for _, _, child in sorted(keyed, key=lambda k: (k[0], k[1]))]
            if ordered == children:
                continue
            for child in ordered:
                element.append(child)  # append moves the existing node
            fixed += 1
    return fixed


def apply_rtl_direction(output_path: Path) -> int:
    """Tag right-to-left paragraphs in a converted DOCX; returns the count.

    pdf2docx reproduces glyph positions but never emits direction markup, so
    Arabic/Hebrew output lands in left-to-right paragraphs: Word then applies
    the bidi algorithm with the wrong base direction and trailing punctuation,
    digits and embedded Latin words jump to the wrong end of the line.

    Only ``w:bidi`` (paragraph base direction) and ``w:rtl`` (run direction)
    are set. Alignment is left alone — pdf2docx derives it from the real glyph
    geometry, so it already matches the page — and section/table direction is
    left alone too, because those reverse column order that pdf2docx has
    already laid out visually.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:  # pragma: no cover - deployment problem
        return 0

    document = Document(str(output_path))
    flipped = 0

    for paragraph in _iter_paragraphs(document):
        text = paragraph.text
        if not text:
            continue
        letters = sum(1 for ch in text if ch.isalpha())
        if not letters:
            continue
        if len(_RTL_CHARS.findall(text)) / letters < _RTL_DOMINANCE:
            continue

        pPr = paragraph._p.get_or_add_pPr()
        if pPr.find(qn("w:bidi")) is None:
            pPr.insert_element_before(OxmlElement("w:bidi"), *_PPR_AFTER_BIDI)
        flipped += 1

        # Mark only the runs that actually hold RTL text: Latin words and
        # numbers inside an RTL paragraph must stay LTR for bidi to work.
        for run in paragraph.runs:
            if not _RTL_CHARS.search(run.text):
                continue
            rPr = run._r.get_or_add_rPr()
            if rPr.find(qn("w:rtl")) is None:
                rPr.insert_element_before(OxmlElement("w:rtl"), *_RPR_AFTER_RTL)

    reordered = _normalize_property_order(document)
    if flipped or reordered:
        document.save(str(output_path))
    if reordered:
        logger.debug("docx_property_order_normalized", elements=reordered)
    return flipped


def _fallback_pdf_to_docx(
    input_path: Path,
    output_path: Path,
    *,
    first_page: int | None,
    last_page: int | None,
) -> None:
    """Last-resort PDF→DOCX via PyMuPDF text extraction.

    Produces paragraphs from each page's text blocks and embeds a page
    render for pages with no extractable text (scans). Loses complex layout
    — acceptable only because the alternative is failing the file outright.
    """
    import fitz
    from docx import Document
    from docx.shared import Inches

    document = Document()
    with fitz.open(str(input_path)) as pdf:
        start = (first_page - 1) if first_page else 0
        stop = min(last_page or pdf.page_count, pdf.page_count)
        for number in range(start, stop):
            page = pdf[number]
            wrote_text = False
            for block in page.get_text("blocks", sort=True):
                content = block[4].strip()
                if content and block[6] == 0:  # text blocks only, not images
                    document.add_paragraph(content)
                    wrote_text = True
            if not wrote_text:
                pixmap = page.get_pixmap(dpi=150)
                document.add_picture(
                    io.BytesIO(pixmap.tobytes("png")),
                    width=Inches(min(6.5, pixmap.width / 150)),
                )
            if number + 1 < stop:
                document.add_page_break()
    document.save(str(output_path))


def pdf_to_docx(
    input_path: Path,
    output_path: Path,
    *,
    first_page: int | None = None,
    last_page: int | None = None,
    page_workers: int = 0,
    display_name: str | None = None,
) -> Path:
    """Convert one PDF to DOCX; returns ``output_path``.

    ``page_workers`` > 1 enables pdf2docx's page-level multi-processing —
    only pass it when :func:`page_parallel_workers` says it is worthwhile.
    ``display_name`` is the user-facing name for error messages (on-disk
    names are storage UUIDs the user has never seen).
    """
    try:
        from pdf2docx import Converter
    except ImportError as exc:  # pragma: no cover - deployment problem
        raise ProcessingError(
            "The PDF-to-Word converter is not installed on this server."
        ) from exc

    name = display_name or input_path.name

    # pdf2docx logs noisily at INFO; keep worker logs clean.
    logging.getLogger("pdf2docx").setLevel(logging.WARNING)

    parallel_kwargs = (
        {"multi_processing": True, "cpu_count": page_workers}
        if page_workers > 1
        else {}
    )

    try:
        converter = Converter(str(input_path))
    except Exception as exc:
        raise ProcessingError(
            f"'{name}' could not be opened as a PDF — the file appears to be "
            "damaged."
        ) from exc
    try:
        converter.convert(
            str(output_path),
            start=(first_page - 1) if first_page else 0,
            end=last_page,  # None = to the end; pdf2docx end is exclusive-ish
            **parallel_kwargs,
        )
    except Exception as exc:
        # Layout analysis gave up on some page. Degrade to the simple
        # extractor rather than failing a file every PDF viewer can read.
        logger.warning(
            "pdf2docx_failed_using_fallback", input=name, error=str(exc)
        )
        try:
            _fallback_pdf_to_docx(
                input_path, output_path, first_page=first_page, last_page=last_page
            )
        except Exception:
            raise ProcessingError(
                f"Could not convert '{name}' to Word. The file may be damaged "
                f"or use unsupported features. (Converter said: {exc})"
            ) from exc
    finally:
        converter.close()

    if not output_path.is_file() or output_path.stat().st_size == 0:
        raise ProcessingError(
            f"Converting '{name}' produced an empty document."
        )

    # Direction markup is a finishing pass over a file that already converted:
    # a failure here must downgrade to LTR output, never fail the job.
    try:
        flipped = apply_rtl_direction(output_path)
        if flipped:
            logger.info("rtl_paragraphs_marked", input=name, paragraphs=flipped)
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning("rtl_pass_failed", input=name, error=str(exc))

    return output_path
